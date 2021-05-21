from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests, docx

numner = 6925

data = {"XesAnalyticsGid": "DD071B71-EE1B-4494-8CBB-2F0FC3D3896D",
        "Hm_lvt_f3c0713d8cd3adc9dd26728c8ecd8e3c": 1621564180,
        "SL_GWPT_Show_Hide_tmp": 1,
        "SL_wptGlobTipTmp": 1,
        "Hm_lvt_9047c8b1678e5280b8806d3a9f9d439f": "1621564024, 1621571778",
        "Hm_lpvt_9047c8b1678e5280b8806d3a9f9d439f": 1621571832}
headers = {"method": "GET",
           "authority": "www.xiaohoucode.com",
           "scheme": "https",
           "path": "/api/core/choice/answer/studentGet?questionId=6929&classId=79c0860c0d3d4ab5aa0004f92d7ee4b9&lessonId=1794001f952000163e0b7f06000cb001",
           "authorization": "eyJhbGciOiJIUzI1NiIsInppcCI6IkdaSVAifQ.H4sIAAAAAAAAAFWO32rCMBTG3yXXDpKTNE29G1pBEAXXgXojaRI1WBtsU7YpAx9HvNtLjT3GEr3y8vy-P-c7I9u2qI8-rXQ71ymnbb1FPWSlR33CgSScAaM91PpOm9qPdTBT0JhIrLKMCaZwJpQCAEG5LlNeliLkGyOrqTyY4P79uf1dL5G5Kt7z2SRfvxXvw3xaBKqs_xo4HRVMcABeVvcvhaxOm9KMjnW-Wu59YZeLlzThr-uPuC86CGMZUEhii2wiAUJDCUuJSGJV15qmfqx4Ur7_AeCjf-D3AAAA.UUvZ1-mwgOFyJBLjvmgmg6utLY_muMDza89bOHl72p4",
           "sec-ch-ua-mobile": "?0",
           }


def word(title, practiceTotal, start_number):
    doc = docx.Document()  # 创建
    # 添加标题title
    doc.add_heading(title, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i in range(start_number, start_number + practiceTotal):
        url = 'https://www.xiaohoucode.com/api/core/choice/answer/studentGet?questionId={}&classId=79c0860c0d3d4ab5aa0004f92d7ee4b9&lessonId=1794001f952000163e0b7f06000cb001'.format(
            i)
        res = requests.get(url, headers=headers).json()
        topic = res["data"]["content"]  # 题目

        optionItems = {}
        for i in res["data"]["optionItems"]:
            optionItems[i["no"]] = i["content"]
            if i['isCorrect'] == True:
                optionItems["正确答案"] = i['no']  # 选项和答案
            print(optionItems)

        # 添加题目段落
        p1 = doc.add_paragraph(style='List Number')
        p1.add_run(topic)
        # print('标题天剑成功')
        # 选项

        for key, value in optionItems.items():
            doc.add_paragraph().add_run(key + ":\t\t\t" + value)
            # print('选项添加成功')
    doc.save(title + ".docx")


url = 'https://www.xiaohoucode.com/api/core/lesson/s-practices?classId=79c0860c0d3d4ab5aa0004f92d7ee4b9&curriculumNo=5'
res = requests.get(url, headers=headers, data=data).json()
for i in res['data']:
    title = i['bundleInfo']['bundleName']
    practiceTotal = i['practiceTotal']
    start_number = i['practices'][0]['projectId']
    print(title)
    print(practiceTotal)
    print(start_number)
    word(title, practiceTotal, start_number)
